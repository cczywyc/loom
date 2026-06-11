import re
from pathlib import Path

import docx

from loom.models import ParsedDocument


def _heading_level(style_name: str) -> int | None:
    """Word 标题样式 → markdown 标题级数（Title=#，Heading N=N+1，衔接 loom ##~###### 节约定）。"""
    s = (style_name or "").strip().lower()
    if s == "title":
        return 1
    m = re.fullmatch(r"heading (\d+)", s)
    return min(int(m.group(1)) + 1, 6) if m else None


def parse_docx(path: Path, assets_dir: Path) -> ParsedDocument:
    """python-docx 逐段抽取；Word 标题样式转成 markdown 标题，保留层级。"""
    document = docx.Document(str(path))
    lines: list[str] = []
    title: str | None = None
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        level = _heading_level(para.style.name if para.style else "")
        if level is not None:
            if title is None:
                title = text
            lines.append("#" * level + " " + text)
        else:
            lines.append(text)
    metadata = {"title": title} if title else {}
    return ParsedDocument(
        source_path=str(path), text="\n\n".join(lines), metadata=metadata, assets=[]
    )
