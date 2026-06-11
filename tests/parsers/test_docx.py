from docx import Document

from loom.parsers import parse_file


def test_docx_parses_with_heading_levels(tmp_path):
    f = tmp_path / "doc.docx"
    d = Document()
    d.add_heading("研究笔记", level=1)
    d.add_paragraph("第一段正文内容。")
    d.add_paragraph("第二段正文内容。")
    d.save(str(f))
    doc = parse_file(f, assets_dir=tmp_path / "assets", wrap=False)
    assert "## 研究笔记" in doc.text  # 一级 Word 标题 → markdown 二级（衔接 loom 节约定）
    assert "第一段正文内容。" in doc.text and "第二段正文内容。" in doc.text
